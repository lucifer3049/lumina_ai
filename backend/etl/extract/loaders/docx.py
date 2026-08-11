"""Word (.docx) loader（08 §3：python-docx）。

docx 與 PDF 最大的差別：**這裡的標題是樣式，不是猜的**。``Heading 2`` 就是第二層，
沒有啟發式，heading_path 可以完全信任。兩種來源產出同一個欄位但可信度不同——1C 評估
檢索品質時要記得這件事。

**段落與表格必須依文件順序交錯讀取**：``document.paragraphs`` 與 ``document.tables``
是兩份各自有序的清單，先讀完段落再讀表格的話，所有表格都會被排到文件末尾，而它們的
heading_path 會全部指向最後一節。因此這裡直接走 XML body 的子元素。

（本模組名為 ``docx`` 而其相依套件也叫 ``docx``：Python 3 的 import 是絕對的，
``import docx`` 取到的是第三方套件，不會遞迴到自己。）
"""

from __future__ import annotations

import io
import re
from typing import Any

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from core.exceptions import ExtractionFailedError
from etl.extract.loaders.base import HeadingStack, LoadResult
from etl.extract.markdown import rows_to_markdown_table
from etl.extract.model import Block, BlockMeta, BlockType

# "Heading 1"、"heading 2"…。Word 的中文介面樣式名仍是英文（樣式的 ID 不隨語言變）。
_HEADING_STYLE = re.compile(r"^heading\s+(\d+)$", re.IGNORECASE)


def load(content: bytes) -> LoadResult:
    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as exc:
        raise ExtractionFailedError(
            "docx 開啟失敗",
            cause=type(exc).__name__,
        ) from exc

    blocks: list[Block] = []
    headings = HeadingStack()
    table_count = 0

    try:
        body = document.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                _append_paragraph(Paragraph(child, document), blocks, headings)
            elif child.tag == qn("w:tbl"):
                _append_table(Table(child, document), blocks, headings)
                table_count += 1
    except ExtractionFailedError:
        raise
    except Exception as exc:
        raise ExtractionFailedError(
            "docx 內容解析失敗",
            cause=type(exc).__name__,
        ) from exc

    doc_meta: dict[str, Any] = {"table_count": table_count}
    return blocks, doc_meta


def _append_paragraph(paragraph: Paragraph, blocks: list[Block], headings: HeadingStack) -> None:
    text = paragraph.text.strip()
    if not text:
        # 空段落是排版產物（用 Enter 拉開間距），不是內容。
        return

    level = _heading_level(paragraph)
    if level is None:
        blocks.append(
            Block(
                type=BlockType.PARAGRAPH,
                text=text,
                # docx 沒有頁的概念：分頁由排版引擎在開啟時決定，檔案裡沒有這個資訊。
                meta=BlockMeta(order=len(blocks), heading_path=headings.path()),
            )
        )
        return

    headings.push(level, text)
    blocks.append(
        Block(
            type=BlockType.HEADING,
            text=text,
            meta=BlockMeta(order=len(blocks), heading_path=headings.path()[:-1]),
        )
    )


def _append_table(table: Table, blocks: list[Block], headings: HeadingStack) -> None:
    """整張表格成為一個 block，內容寫成 **GFM 表格**，表頭留著。

    表頭決定每一欄的意義。攤平成純文字後，「2025 3.2 通過」這種列對 LLM 毫無意義；
    1B-5 的「表格不拆散」與 1C 的檢索都需要它。表格因此也不在這裡切——切塊是下游的
    事，這裡只保證它是一個完整、可辨識型別的 block。

    格式在 loader 就寫成 Markdown（而不是留給渲染器）是因為**只有這裡看得到儲存格
    的行列邊界**：一旦攤平成字串，含分隔符的儲存格就無法還原成正確的欄位。
    """
    rows = [[cell.text for cell in row.cells] for row in table.rows]
    text = rows_to_markdown_table([row for row in rows if any(cell.strip() for cell in row)])
    if not text.strip():
        return

    blocks.append(
        Block(
            type=BlockType.TABLE,
            text=text,
            meta=BlockMeta(order=len(blocks), heading_path=headings.path()),
        )
    )


def _heading_level(paragraph: Paragraph) -> int | None:
    """樣式名 → 標題層級；不是標題就回 None。

    ``style`` 可能是 None（樣式被刪除的畸形檔）——那時當內文處理，不是失敗：
    段落的文字仍然是有效內容，丟掉它比把層級判錯糟得多。
    """
    style = paragraph.style
    name = getattr(style, "name", None)
    if not name:
        return None
    match = _HEADING_STYLE.match(name)
    return int(match.group(1)) if match else None
