"""驗收：三種 loader → 統一的 `ExtractedDoc`（08 §3、13 §3 工作包 1B-4）。

**loader 的職責邊界只有一件事**：來源 → 統一中間格式。清洗（頁首頁尾偵測、亂碼
丟棄）屬 Clean、切塊屬 Chunk，兩者都在 1B-5——混進 loader 的話，每加一種來源就要
重寫一次同樣的清洗規則，而 08 §1 的「下游完全來源無關」也就不成立了。

因此本檔驗的是**結構保真**，不是內容品質：

1. **block 型別分得出來**（paragraph / heading / table）。分不出來的話，1B-5 的
   「標題邊界優先、表格不拆散」就沒有依據可用——chunker 只會看到一堆等價的文字。
2. **標題階層（heading_path）**：chunk 的 meta 要靠它回答「這段話出自哪一節」，
   而那是引用能不能指到正確位置的前提。
3. **頁碼**：PDF 的引用要標頁數（09 §3.2 的 citations event 有 page 欄位）。
4. **順序（order）**：block 的先後就是文件的閱讀順序。亂了不會有錯誤，只會讓答案
   引用到語意上不相鄰的段落。

測試用的樣本檔在測試裡**即時產生**而不是放進版控：二進位樣本進 repo 之後沒有人
知道它們裡面有什麼，改壞了也看不出來；用程式產生的話，「這份 PDF 有兩頁、第二頁
有一個標題」是讀得到的。
"""

from __future__ import annotations

import io
import zipfile

import pytest

from core.exceptions import ExtractionFailedError
from core.media_types import MARKDOWN, XLSX
from etl.extract import extract
from etl.extract.loaders.xlsx import ROWS_PER_BLOCK
from etl.extract.model import BlockType, ExtractedDoc
from services.knowledge.uploads import DOCX, PDF, TEXT


def _pdf(
    pages: list[list[tuple[str, int]]],
    *,
    outline: list[tuple[int, str, int]] | None = None,
    tables: dict[int, list[list[str]]] | None = None,
) -> bytes:
    """產生 PDF。

    ``pages``：每頁一串 (文字, 字級)。字級大的視為標題（loader 的啟發式）。
    ``outline``：(層級, 標題, 頁碼) 的 PDF 大綱項；有大綱時 loader 應優先採信它。
    ``tables``：頁碼 → 列資料；畫出**有框線**的表格供 loader 偵測。

    用 reportlab（BSD）而不是解析用的那個函式庫產生樣本，除了授權，還有一個測試上的
    理由：產生與解析若共用同一套實作，兩邊的偏差會互相抵銷——樣本寫錯什麼，解析就
    照樣讀回什麼，測試全綠而真實 PDF 全錯。

    中文必須註冊 CID 字型（1B-4 實作時踩過一次）：預設的 base-14 字型沒有中文字形，
    寫進去的字會變成替代符號，抽出來自然對不上——**那是樣本產不出中文，不是 loader
    抽不到**。
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas

    font = "STSong-Light"
    pdfmetrics.registerFont(UnicodeCIDFont(font))

    outline_by_page: dict[int, list[tuple[int, str]]] = {}
    for level, title, page_number in outline or []:
        outline_by_page.setdefault(page_number, []).append((level, title))

    buffer = io.BytesIO()
    _, height = A4
    pdf = canvas.Canvas(buffer, pagesize=A4)

    for page_index, lines in enumerate(pages, start=1):
        y = height - 72.0
        for text, size in lines:
            pdf.setFont(font, size)
            pdf.drawString(72, y, text)
            # 行距 1.3 倍字級是排版慣例。用更大的值（例如兩倍）產生的樣本，每一行看
            # 起來都像另起一段——loader 的分段判定會被測成「永遠不合併」，而那是樣本
            # 不真實，不是實作有問題。
            y -= size * 1.3

        for level, title in outline_by_page.get(page_index, []):
            key = f"p{page_index}-{level}-{title}"
            pdf.bookmarkPage(key)
            # reportlab 的 level 從 0 起算，PDF 大綱的層級由巢狀關係決定。
            pdf.addOutlineEntry(title, key, level=level - 1)

        rows = (tables or {}).get(page_index)
        if rows:
            y -= 24
            cell_width, row_height = 120, 24
            for row_index, row in enumerate(rows):
                for column_index, value in enumerate(row):
                    x = 72 + column_index * cell_width
                    top = y - row_index * row_height
                    pdf.rect(x, top, cell_width, row_height)
                    pdf.setFont(font, 11)
                    pdf.drawString(x + 4, top + 8, value)

        pdf.showPage()

    pdf.save()
    return buffer.getvalue()


def _xlsx(sheets: dict[str, list[list[object]]]) -> bytes:
    """產生 xlsx：工作表名 → 列資料（第一列視為表頭）。"""
    import openpyxl

    workbook = openpyxl.Workbook()
    # 新活頁簿自帶一張空白工作表；留著它會多出一個空節。
    default_sheet = workbook.active
    if default_sheet is not None:
        workbook.remove(default_sheet)
    for title, rows in sheets.items():
        sheet = workbook.create_sheet(title=title)
        for row in rows:
            sheet.append(row)

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _docx(paragraphs: list[tuple[str, str]], tables: list[list[list[str]]] | None = None) -> bytes:
    """產生 docx：段落是 (文字, 樣式名)，樣式 ``Heading 1`` 之類代表標題。"""
    import docx

    document = docx.Document()
    for text, style in paragraphs:
        document.add_paragraph(text, style=style or None)
    for rows in tables or []:
        table = document.add_table(rows=len(rows), cols=len(rows[0]))
        for row_index, row in enumerate(rows):
            for cell_index, value in enumerate(row):
                table.cell(row_index, cell_index).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestPlainText:
    def test_paragraphs_are_split_on_blank_lines(self) -> None:
        """空行分段——純文字唯一可靠的結構訊號。

        不以單一換行分段：那會把「一段話因為視窗寬度而斷行」拆成好幾個 block，
        而 chunker 拿到的每一塊都短到沒有語意。
        """
        content = "第一段第一行\n第一段第二行\n\n第二段\n".encode()

        doc = extract(content, media_type=TEXT)

        assert [block.text for block in doc.blocks] == ["第一段第一行\n第一段第二行", "第二段"]
        assert {block.type for block in doc.blocks} == {BlockType.PARAGRAPH}

    def test_order_is_preserved(self) -> None:
        content = "\n\n".join(f"段落{i}" for i in range(5)).encode()

        doc = extract(content, media_type=TEXT)

        assert [block.meta.order for block in doc.blocks] == [0, 1, 2, 3, 4]

    def test_text_has_no_page_numbers(self) -> None:
        """純文字沒有頁的概念——``page`` 必須是 None，不是 0 或 1。

        填一個假頁碼的話，引用會顯示「第 1 頁」而那個頁不存在；使用者點過去只會
        更困惑。沒有的資訊就要表達成沒有。
        """
        doc = extract(b"only one paragraph", media_type=TEXT)

        assert doc.blocks[0].meta.page is None


class TestPdf:
    def test_text_and_page_numbers_are_extracted(self) -> None:
        content = _pdf([[("第一頁的內容", 11)], [("第二頁的內容", 11)]])

        doc = extract(content, media_type=PDF)

        texts = {block.meta.page: block.text for block in doc.blocks}
        assert "第一頁的內容" in texts[1]
        assert "第二頁的內容" in texts[2]

    def test_page_numbers_are_one_based(self) -> None:
        """頁碼從 1 開始——那是使用者看到的編號。

        解析函式庫的頁索引常是 0-based，直接沿用會讓引用整份文件差一頁，而每一個
        引用看起來都很合理（只是內容對不上）。
        """
        doc = extract(_pdf([[("唯一一頁", 11)]]), media_type=PDF)

        assert {block.meta.page for block in doc.blocks} == {1}

    def test_larger_font_becomes_a_heading(self) -> None:
        """字級明顯大於內文的行視為標題。

        PDF 沒有語意標記，只有視覺屬性——標題偵測必然是啟發式的。用字級而不是粗體
        或位置，是因為字級是三者中最少誤判的：粗體常用於強調，而位置對雙欄排版
        完全失效。判錯的後果是 heading_path 不準（引用指到相鄰的節），不是壞掉。
        """
        content = _pdf([[("第一章 總則", 22), ("本章說明適用範圍。", 11)]])

        doc = extract(content, media_type=PDF)
        headings = [block for block in doc.blocks if block.type is BlockType.HEADING]

        assert [block.text for block in headings] == ["第一章 總則"]

    def test_body_blocks_carry_the_heading_path(self) -> None:
        """內文的 ``heading_path`` 是它上方所有標題的堆疊。

        這是 1B-5 切塊與 1D 引用共同的依據：使用者看到的是「出自〈第一章 總則〉」，
        而不是一個沒有上下文的段落。
        """
        content = _pdf([[("第一章 總則", 22), ("本章說明適用範圍。", 11)]])

        doc = extract(content, media_type=PDF)
        body = next(block for block in doc.blocks if block.type is BlockType.PARAGRAPH)

        assert body.meta.heading_path == ("第一章 總則",)

    def test_page_count_is_reported_in_doc_meta(self) -> None:
        """頁數進 ``doc_meta``——08 §4 的 stats 要它，而且是毒檔上限的判定依據。"""
        doc = extract(_pdf([[("a", 11)], [("b", 11)], [("c", 11)]]), media_type=PDF)

        assert doc.doc_meta["page_count"] == 3

    def test_wrapped_lines_become_one_paragraph(self) -> None:
        """同一段的折行要合回一個 block。

        PDF 裡沒有「段落」，只有一行一行的字。逐行成塊的話，chunk 會是一堆各自
        殘缺的短句，檢索時每一塊都缺少判斷語意所需的上下文。
        """
        content = _pdf([[("第一段的第一行，", 11), ("第一段的第二行。", 11)]])

        doc = extract(content, media_type=PDF)
        paragraphs = [block for block in doc.blocks if block.type is BlockType.PARAGRAPH]

        assert len(paragraphs) == 1
        assert "第一行" in paragraphs[0].text and "第二行" in paragraphs[0].text


class TestPdfOutline:
    """PDF 大綱（書籤）優先於字級啟發式（1B-4c）。

    大綱是作者**明確標記**的階層，不是我們從版面猜出來的。有大綱時還去猜字級，等於
    把已知的正確答案丟掉換一個估計值——而那個估計值在字級一致的文件上必然失敗。
    """

    def test_outline_titles_become_headings_even_at_body_size(self) -> None:
        content = _pdf(
            [[("第一章 總則", 11), ("本章說明適用範圍。", 11)]],
            outline=[(1, "第一章 總則", 1)],
        )

        doc = extract(content, media_type=PDF)
        headings = [block for block in doc.blocks if block.type is BlockType.HEADING]

        assert [block.text for block in headings] == ["第一章 總則"]

    def test_outline_level_drives_the_heading_path(self) -> None:
        """大綱的層級直接決定 heading_path 的深度，不必再由字級排名推。"""
        content = _pdf(
            [
                [("第一章", 11), ("第一節", 11), ("本節內容。", 11)],
            ],
            outline=[(1, "第一章", 1), (2, "第一節", 1)],
        )

        doc = extract(content, media_type=PDF)
        body = next(block for block in doc.blocks if block.type is BlockType.PARAGRAPH)

        assert body.meta.heading_path == ("第一章", "第一節")


class TestPdfTables:
    """表格抽成 TABLE block（1B-4c）。

    使用者說「PDF 切出來超級不準確」的典型症狀就在這裡：一張表被攤成一串沒有欄位
    歸屬的數字。抽成獨立 block 並保留表頭，1B-5 才有辦法「表格不拆散」。
    """

    def test_ruled_table_becomes_a_markdown_table_block(self) -> None:
        content = _pdf([[("規格表", 22)]], tables={1: [["項目", "數值"], ["延遲", "300ms"]]})

        doc = extract(content, media_type=PDF)
        tables = [block for block in doc.blocks if block.type is BlockType.TABLE]

        assert len(tables) == 1
        assert tables[0].text.splitlines()[0] == "| 項目 | 數值 |"
        assert "| 延遲 | 300ms |" in tables[0].text

    def test_table_text_is_not_repeated_as_paragraphs(self) -> None:
        """表格內的文字不得同時以段落形式再出現一次。

        重複的後果不是「多一份」而已：同一份內容會被切成兩種形狀各自嵌入，檢索時
        互相排擠，而回答引用到的是那個沒有欄位結構的版本。
        """
        content = _pdf([[("規格表", 22)]], tables={1: [["項目", "數值"], ["延遲", "300ms"]]})

        doc = extract(content, media_type=PDF)
        paragraphs = [block for block in doc.blocks if block.type is BlockType.PARAGRAPH]

        assert all("300ms" not in block.text for block in paragraphs)

    def test_table_carries_page_and_heading_path(self) -> None:
        """表格也要能被引用——頁碼與所屬章節缺一不可。"""
        content = _pdf(
            [[("規格表", 11)]],
            outline=[(1, "規格表", 1)],
            tables={1: [["項目", "數值"], ["延遲", "300ms"]]},
        )

        doc = extract(content, media_type=PDF)
        table = next(block for block in doc.blocks if block.type is BlockType.TABLE)

        assert table.meta.page == 1
        assert table.meta.heading_path == ("規格表",)


class TestDocx:
    def test_headings_are_detected_from_styles(self) -> None:
        """docx 的標題是**樣式**，不是猜的——這是它與 PDF 最大的差別。

        因此 docx 的 heading_path 可以完全信任，而 PDF 的是啟發式的。兩者產出同一個
        欄位，但可信度不同；1C 評估檢索品質時要記得這件事。
        """
        content = _docx(
            [
                ("第一章 總則", "Heading 1"),
                ("第一節 目的", "Heading 2"),
                ("本節說明立法目的。", ""),
            ]
        )

        doc = extract(content, media_type=DOCX)
        body = next(block for block in doc.blocks if block.type is BlockType.PARAGRAPH)

        assert body.meta.heading_path == ("第一章 總則", "第一節 目的")

    def test_heading_path_pops_when_the_level_goes_back_up(self) -> None:
        """回到上層標題時，下層要從路徑裡拿掉。

        只往裡疊不往外彈的實作，會讓文件後半的每一段都掛著前面所有小節的名字——
        heading_path 變成一條愈來愈長的垃圾，引用顯示出來完全不可讀。
        """
        content = _docx(
            [
                ("第一章", "Heading 1"),
                ("第一節", "Heading 2"),
                ("甲段", ""),
                ("第二章", "Heading 1"),
                ("乙段", ""),
            ]
        )

        doc = extract(content, media_type=DOCX)
        paragraphs = [b for b in doc.blocks if b.type is BlockType.PARAGRAPH]

        assert paragraphs[0].meta.heading_path == ("第一章", "第一節")
        assert paragraphs[1].meta.heading_path == ("第二章",)

    def test_tables_become_table_blocks_with_the_header_kept(self) -> None:
        """表格是獨立的 block 型別，且表頭要留著。

        表頭決定每一欄的意義。把表格攤平成純文字之後，「2025 / 3.2 / 通過」這種列
        對 LLM 毫無意義；1B-5 的「表格不拆散」與 1C 的檢索都需要它。
        """
        content = _docx(
            [("規格表", "Heading 1")],
            tables=[[["項目", "數值"], ["延遲", "300ms"]]],
        )

        doc = extract(content, media_type=DOCX)
        tables = [block for block in doc.blocks if block.type is BlockType.TABLE]

        assert len(tables) == 1
        assert "項目" in tables[0].text and "延遲" in tables[0].text
        assert tables[0].meta.heading_path == ("規格表",)

    def test_empty_paragraphs_are_dropped(self) -> None:
        """空段落不產生 block——它們是排版產物，不是內容。"""
        content = _docx([("有內容", ""), ("", ""), ("   ", ""), ("也有內容", "")])

        doc = extract(content, media_type=DOCX)

        assert [block.text for block in doc.blocks] == ["有內容", "也有內容"]


class TestXlsx:
    """試算表（08 §3：每 sheet 一節、表頭隨列窗重複；1B-4b）。

    試算表與文件的差別在於**它沒有敘事**：一張表就是一大片格子，沒有段落、沒有頁。
    所以結構全靠兩件事撐住——工作表名（它就是這一節的標題）與表頭（它決定每一欄的
    意義）。少了任一個，切出來的塊就是一串沒有歸屬的數字。
    """

    def test_each_sheet_becomes_a_section(self) -> None:
        content = _xlsx({"營收": [["年", "金額"], ["2025", "100"]], "成本": [["年", "金額"]]})

        doc = extract(content, media_type=XLSX)
        headings = [block.text for block in doc.blocks if block.type is BlockType.HEADING]

        assert headings == ["營收", "成本"]

    def test_rows_become_a_table_block_under_the_sheet_heading(self) -> None:
        content = _xlsx({"營收": [["年", "金額"], ["2025", "100"]]})

        doc = extract(content, media_type=XLSX)
        table = next(block for block in doc.blocks if block.type is BlockType.TABLE)

        assert table.meta.heading_path == ("營收",)
        assert table.text.splitlines()[0] == "| 年 | 金額 |"
        assert "| 2025 | 100 |" in table.text

    def test_large_sheet_is_split_into_windows_that_each_repeat_the_header(self) -> None:
        """大表按列窗切開，**每一塊都自帶表頭**。

        表頭只放在第一塊的話，第二塊之後的每一列都會失去欄位意義——而檢索命中的
        往往正是後面那些塊（它們才是資料）。
        """
        rows: list[list[object]] = [["年", "金額"]]
        rows += [[str(2000 + i), str(i)] for i in range(ROWS_PER_BLOCK + 5)]

        doc = extract(_xlsx({"營收": rows}), media_type=XLSX)
        tables = [block for block in doc.blocks if block.type is BlockType.TABLE]

        assert len(tables) == 2
        assert all(table.text.splitlines()[0] == "| 年 | 金額 |" for table in tables)
        # 資料列不得因為切窗而遺失或重複。
        data_lines = [
            line
            for table in tables
            for line in table.text.splitlines()[2:]  # 跳過表頭與分隔列
        ]
        assert len(data_lines) == ROWS_PER_BLOCK + 5

    def test_empty_sheets_are_skipped(self) -> None:
        """完全空的工作表不產生任何 block——它是活頁簿的殘留，不是內容。"""
        doc = extract(_xlsx({"有資料": [["a"], ["1"]], "空的": []}), media_type=XLSX)

        assert [block.text for block in doc.blocks if block.type is BlockType.HEADING] == ["有資料"]

    def test_sheet_count_is_reported_in_doc_meta(self) -> None:
        doc = extract(_xlsx({"一": [["a"]], "二": [["b"]]}), media_type=XLSX)

        assert doc.doc_meta["sheet_count"] == 2

    def test_spreadsheets_have_no_page_numbers(self) -> None:
        """試算表沒有頁——``page`` 必須是 None（同純文字的理由）。"""
        doc = extract(_xlsx({"一": [["a"], ["1"]]}), media_type=XLSX)

        assert all(block.meta.page is None for block in doc.blocks)


class TestMarkdown:
    """Markdown 來源（1B-4b）。

    Markdown 是**唯一結構已經明說的文字來源**：標題就是標題，表格就是表格，不必
    像 PDF 那樣猜。因此這裡的錯誤方式只有一種——把已經標好的結構弄丟。
    """

    def test_headings_build_the_heading_path(self) -> None:
        content = "# 第一章\n\n## 第一節\n\n本節內容。\n".encode()

        doc = extract(content, media_type=MARKDOWN)
        body = next(block for block in doc.blocks if block.type is BlockType.PARAGRAPH)

        assert body.meta.heading_path == ("第一章", "第一節")

    def test_heading_text_excludes_the_hashes(self) -> None:
        """``#`` 是標記不是內容——留著它，引用顯示與 heading_path 都會帶著井字號。"""
        doc = extract(b"## Section title\n", media_type=MARKDOWN)

        assert doc.blocks[0].text == "Section title"

    def test_gfm_table_stays_a_table_block(self) -> None:
        content = "| 項目 | 數值 |\n| --- | --- |\n| 延遲 | 300ms |\n".encode()

        doc = extract(content, media_type=MARKDOWN)
        table = next(block for block in doc.blocks if block.type is BlockType.TABLE)

        assert "| 延遲 | 300ms |" in table.text

    def test_fenced_code_becomes_a_code_block(self) -> None:
        """程式碼要標成 CODE：1B-5 不該在函式中間切開，而切塊靠的是 block 型別。"""
        content = b"```python\nprint('hi')\n```\n"

        doc = extract(content, media_type=MARKDOWN)

        assert [block.type for block in doc.blocks] == [BlockType.CODE]
        assert "print('hi')" in doc.blocks[0].text

    def test_a_list_stays_one_block(self) -> None:
        """整份清單是一個 block。

        逐項成塊的話，「以下三種情況適用：」與它的三個項目會被切散，而每一項單獨
        看都不知道在說什麼。
        """
        content = "- 第一項\n- 第二項\n- 第三項\n".encode()

        doc = extract(content, media_type=MARKDOWN)

        assert len(doc.blocks) == 1
        assert "第三項" in doc.blocks[0].text

    def test_markdown_has_no_page_numbers(self) -> None:
        doc = extract(b"just text\n", media_type=MARKDOWN)

        assert doc.blocks[0].meta.page is None


class TestUnsupportedAndBroken:
    def test_unknown_media_type_raises(self) -> None:
        """沒有 loader 的型別要明確失敗。

        08 §1：新增來源 = 新增一個 loader。回空的 `ExtractedDoc` 會讓文件走完整條
        ETL 並停在 ready，而它一個 chunk 都沒有——使用者問問題時只會覺得「這份文件
        好像沒有被讀進去」。
        """
        with pytest.raises(ExtractionFailedError):
            extract(b"%PDF-1.7\n", media_type="application/x-tar")

    def test_corrupted_pdf_raises_extraction_failed(self) -> None:
        """壞掉的檔案 → `ExtractionFailedError`，不是函式庫自己的例外。

        pdfplumber / python-docx 的例外冒到上層之後，08 §6 的重試判定就得認得第三方
        的例外型別——而那會隨版本改變。轉成自家例外，重試與 DLQ 的規則才有穩定依據。
        """
        with pytest.raises(ExtractionFailedError):
            extract("%PDF-1.7\n這不是一個真的 PDF".encode(), media_type=PDF)

    def test_docx_without_document_xml_raises(self) -> None:
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w") as archive:
            archive.writestr("hello.txt", "not a docx")

        with pytest.raises(ExtractionFailedError):
            extract(buffer.getvalue(), media_type=DOCX)


class TestExtractedDocShape:
    def test_doc_meta_records_the_source_media_type(self) -> None:
        """`doc_meta` 要記下它是從什麼格式來的（08 §3）。

        下游（Clean、Chunk、評估）會依格式調整策略——例如 PDF 的頁首頁尾偵測對
        純文字沒有意義。少了這個欄位，那些判斷只能靠猜或再看一次原檔。
        """
        doc = extract(b"hello", media_type=TEXT)

        assert isinstance(doc, ExtractedDoc)
        assert doc.doc_meta["media_type"] == TEXT

    def test_block_count_is_recorded(self) -> None:
        """08 §4 的 stats：block 數是「丟棄率」的分母，1B-5 要用。"""
        doc = extract("一\n\n二\n\n三".encode(), media_type=TEXT)

        assert doc.doc_meta["block_count"] == 3
